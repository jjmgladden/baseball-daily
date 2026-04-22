#!/usr/bin/env python
"""
Builds docs/Daily-Email-Setup-Guide.docx — a comprehensive reference for
the morning-email feature: origin, decisions, architecture, baseball setup,
pickleball setup, GitHub vs local, recipient management, custom domain,
troubleshooting.

One-shot build. Re-run to regenerate the docx from this source.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- helpers ----------

NAVY       = RGBColor(0x0E, 0x18, 0x21)
CARDS_RED  = RGBColor(0xC4, 0x1E, 0x3A)
NATS_RED   = RGBColor(0xAB, 0x00, 0x03)
MUTED      = RGBColor(0x55, 0x5B, 0x63)
GOLD       = RGBColor(0xB8, 0x88, 0x00)
CODE_BG    = "F3F4F6"
ACCENT_BG  = "E8F1FA"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_rich_para(doc, segments):
    """segments = list of (text, {bold, italic, color, code}). code=True -> monospace + light bg not rendered inline; use add_code_block for blocks."""
    p = doc.add_paragraph()
    for seg in segments:
        text, opts = seg
        run = p.add_run(text)
        run.bold = opts.get('bold', False)
        run.italic = opts.get('italic', False)
        if opts.get('color'):
            run.font.color.rgb = opts['color']
        if opts.get('code'):
            run.font.name = 'Consolas'
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), 'Consolas')
            rFonts.set(qn('w:hAnsi'), 'Consolas')
            run.font.size = Pt(10)
        if opts.get('size'):
            run.font.size = Pt(opts['size'])
    return p


def add_bullet(doc, text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    doc.add_paragraph(text, style=style)


def add_numbered(doc, text, level=0):
    style = 'List Number' if level == 0 else 'List Number 2'
    doc.add_paragraph(text, style=style)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    # shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), CODE_BG)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    return p


def add_callout(doc, title, body, color=ACCENT_BG):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, color)
    cell.text = ''
    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    p_body = cell.add_paragraph(body)
    for run in p_body.runs:
        run.font.size = Pt(10)
    doc.add_paragraph()  # spacing


def add_table(doc, headers, rows, col_widths_in=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        cells = t.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ---------- build ----------

def build():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ===== Title =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Daily Morning Email — Complete Setup & Design Record")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = CARDS_RED

    sub = doc.add_paragraph()
    r = sub.add_run("Ozark Joe's Baseball Daily  ·  and the future Pickleball sibling project")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    r = meta.add_run("Version 1  ·  2026-04-22  ·  Source: chat session 2026-04-22  ·  Author: Claude + John Gladden")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    doc.add_paragraph()

    # ===== Part 1: Origin =====
    add_heading(doc, "1.  Origin — The Prompt That Started It", level=1, color=NAVY)

    add_para(doc, "This feature began with the following message from the project owner (John Gladden) in the 2026-04-22 chat session:", italic=True, color=MUTED)

    add_callout(
        doc,
        "Owner prompt — 2026-04-22",
        "\"my thinking is that every morning, after the daily ingestion, you would send him and me an email with the links so that he can just open the email and press the link for the latest. is that workable??\"",
        color=ACCENT_BG,
    )

    add_para(doc, "Context:", bold=True)
    add_bullet(doc, "\"Him\" = the owner's brother in Virginia — the project's target reader alongside the owner himself.")
    add_bullet(doc, "The full site already lives at https://jjmgladden.github.io/baseball-daily/.")
    add_bullet(doc, "Daily ingestion already runs autonomously at 07:00 UTC (3 AM EDT / 2 AM EST) via GitHub Actions.")
    add_bullet(doc, "The brother's workflow is email-centric — he opens email in the morning; asking him to bookmark a URL or install a PWA is friction.")
    add_para(doc, "")

    # ===== Part 2: Decisions =====
    add_heading(doc, "2.  Four Decisions That Shaped the Build", level=1, color=NAVY)
    add_para(doc, "Before writing code, four design questions were raised and the owner answered each. These decisions are the spec:", italic=True, color=MUTED)

    add_table(
        doc,
        ["#", "Question raised", "Owner's decision", "Implication"],
        [
            ["1", "Which email provider?",
             "Resend.com",
             "Developer-focused API, 3,000 emails/month free tier, ~5 lines of code."],
            ["2", "Who receives it? Extensible?",
             "Owner + brother initially; extensible via a simple list.",
             "Recipient list stored as a comma-separated GitHub Secret — no code change to add/remove people."],
            ["3", "Sender domain — default or custom?",
             "Start with Resend's default (onboarding@resend.dev); document the custom-domain path for later.",
             "Zero-cost launch; ~$10/yr upgrade path when branding matters."],
            ["4", "Email format — link only, rich preview, or full content?",
             "Rich preview + link (best of both).",
             "Email shows Cardinals/Nationals results, On This Day, stats footer, and a prominent button to the full site."],
        ],
        col_widths_in=[0.3, 1.7, 2.0, 2.5],
    )

    # ===== Part 3: How it works =====
    add_heading(doc, "3.  How It Works — End-to-End Architecture", level=1, color=NAVY)

    add_heading(doc, "3.1  Daily flow", level=2)
    add_code_block(doc,
        "07:00 UTC  (3 AM EDT / 2 AM EST)\n"
        "     │\n"
        "     ▼\n"
        "GitHub Actions fires .github/workflows/daily.yml\n"
        "     │\n"
        "     ├─[1]  Checkout repo\n"
        "     ├─[2]  Setup Node 20\n"
        "     ├─[3]  Run ingestion/fetch-daily.js\n"
        "     │         → MLB Stats API, YouTube API, On-This-Day seed\n"
        "     │         → writes data/snapshots/YYYY-MM-DD.json + latest.json\n"
        "     ├─[4]  Commit + push snapshot (bot identity)\n"
        "     ├─[5]  GitHub Pages auto-rebuilds the site (1-2 min)\n"
        "     └─[6]  Run ingestion/send-email.js              ◄── NEW\n"
        "              → reads latest.json\n"
        "              → builds HTML + plain-text + subject\n"
        "              → POST https://api.resend.com/emails\n"
        "              → Resend delivers to every address in EMAIL_RECIPIENTS\n"
    )

    add_heading(doc, "3.2  What's in the email", level=2)
    add_bullet(doc, "Subject leads with Cardinals result when available, e.g. \"⚾ Cardinals win 5-3 vs Miami Marlins — Tuesday, April 21, 2026\".")
    add_bullet(doc, "Cardinals pin card (red stripe, win/loss indicator, score, highlight count).")
    add_bullet(doc, "Nationals pin card (Nats red stripe, same structure).")
    add_bullet(doc, "On This Day — up to 2 events.")
    add_bullet(doc, "Big red CTA button — \"Open the full report →\" linking to the site.")
    add_bullet(doc, "Stats footer — games-yesterday · trades · league-wide IL count.")
    add_bullet(doc, "Plain-text fallback for clients that block HTML.")

    add_heading(doc, "3.3  Graceful degradation", level=2)
    add_para(doc, "The email step is a silent no-op if credentials are missing. This means:", bold=False)
    add_bullet(doc, "If RESEND_API_KEY is not configured → ingestion still runs, snapshot still commits, email is skipped with one log line.")
    add_bullet(doc, "If EMAIL_RECIPIENTS is empty → same behavior.")
    add_bullet(doc, "If Resend API returns an error → the email step fails loudly (exit 1), but the snapshot was already committed and the site already updated. No partial state.")
    add_callout(
        doc,
        "Why this matters",
        "Enabling or disabling the email feature is just a GitHub Secret change — no code edits, no YAML edits, no downtime. Delete RESEND_API_KEY to pause; re-add to resume.",
    )

    # ===== Part 4: Baseball setup =====
    add_heading(doc, "4.  Baseball — Complete Setup", level=1, color=CARDS_RED)
    add_para(doc, "One-time setup, ~5 minutes. Steps 1-4 are required to receive email; step 5 is optional.", italic=True, color=MUTED)

    add_heading(doc, "Step 1 — Create a Resend account", level=2)
    add_numbered(doc, "Visit https://resend.com/signup.")
    add_numbered(doc, "Sign up with any email you want as the account owner (this doesn't have to be a recipient).")
    add_numbered(doc, "Verify the email Resend sends.")

    add_heading(doc, "Step 2 — Create an API key", level=2)
    add_numbered(doc, "In the Resend dashboard → left sidebar → API Keys → + Create API Key.")
    add_numbered(doc, "Name: \"Ozark Joe's Baseball Daily\".")
    add_numbered(doc, "Permission: Sending access (not full access).")
    add_numbered(doc, "Click Add, then copy the re_... key immediately — Resend will never show it again. Paste it into Notepad temporarily.")

    add_heading(doc, "Step 3 — Add two GitHub Secrets (in the BASEBALL repo)", level=2)
    add_para(doc, "Go to:", bold=False)
    add_code_block(doc, "https://github.com/jjmgladden/baseball-daily/settings/secrets/actions")
    add_para(doc, "Click \"New repository secret\" twice:")

    add_table(
        doc,
        ["Secret Name (case-sensitive)", "Value"],
        [
            ["RESEND_API_KEY",
             "Paste the re_... key from Step 2."],
            ["EMAIL_RECIPIENTS",
             "Comma-separated addresses, no spaces needed.\nExample:\njjmgladden@gmail.com,brother@gmail.com"],
        ],
        col_widths_in=[2.2, 4.3],
    )

    add_heading(doc, "Step 4 — Test it", level=2)
    add_para(doc, "Manually trigger the daily workflow without waiting for 3 AM:")
    add_code_block(doc, "gh workflow run daily.yml --repo jjmgladden/baseball-daily")
    add_para(doc, "Or visit the Actions tab in GitHub and click Run workflow.")
    add_para(doc, "Verify after ~90 seconds:")
    add_bullet(doc, "Email arrives in your inbox (check spam on first run).")
    add_bullet(doc, "Workflow log ends with: [send-email] Sent. Resend id: <uuid>")
    add_bullet(doc, "Resend dashboard → Emails tab shows the delivery.")

    add_heading(doc, "Step 5 (optional) — Custom sender domain", level=2)
    add_para(doc, "Default: Ozark Joe's Baseball Daily <onboarding@resend.dev>. Works fine but Resend appends a small disclaimer footer.")
    add_para(doc, "To send from daily@ozarkjoe.com or similar:")
    add_numbered(doc, "Register a domain (~$10/yr at Cloudflare Registrar — at-cost pricing, free DNS).")
    add_numbered(doc, "Resend dashboard → Domains → + Add Domain → enter the domain.")
    add_numbered(doc, "Copy the 3-4 DNS records Resend displays (MX, SPF TXT, DKIM TXT, optional DMARC TXT).")
    add_numbered(doc, "Paste them into Cloudflare's DNS panel for that domain. For Cloudflare specifically: make sure proxy status is \"DNS only\" (grey cloud), not proxied.")
    add_numbered(doc, "Click Verify in Resend. Usually completes within 10-30 min; DNS propagation can take up to 24 hours.")
    add_numbered(doc, "Add a third GitHub Secret: EMAIL_FROM = Ozark Joe's Baseball Daily <daily@ozarkjoe.com>")
    add_para(doc, "No code change. Next workflow run sends from the custom address.")

    # ===== Part 5: Pickleball setup =====
    add_heading(doc, "5.  Pickleball — Future Setup (Mirror of Baseball)", level=1, color=GOLD)
    add_para(doc, "When the sibling Pickleball project is built, it will use the same email pattern with one Resend account shared between the two projects. Key differences from baseball are called out below.", italic=True, color=MUTED)

    add_heading(doc, "5.1  Prerequisites — already done from baseball", level=2)
    add_bullet(doc, "Resend account exists — reuse it. Do NOT create a second account.")
    add_bullet(doc, "RESEND_API_KEY exists — the same re_... value works for both projects.")

    add_heading(doc, "5.2  What pickleball needs, separately", level=2)
    add_para(doc, "The Pickleball GitHub repo will have its own Secrets. Secrets do not cross repos.")

    add_table(
        doc,
        ["Secret", "Pickleball action", "Same as baseball?"],
        [
            ["RESEND_API_KEY",
             "Paste the same re_... value into the pickleball repo's Secrets.",
             "Yes — one Resend key covers both projects."],
            ["EMAIL_RECIPIENTS",
             "Set pickleball's recipient list. May be the same people, or different (e.g. brother only cares about baseball).",
             "Usually different. See §7 for sync options."],
            ["EMAIL_FROM (optional)",
             "e.g. pickleball@ozarkjoe.com if a custom subdomain is used.",
             "Free to differ — branded per project."],
        ],
        col_widths_in=[2.0, 3.0, 1.5],
    )

    add_heading(doc, "5.3  Code re-use", level=2)
    add_para(doc, "The pickleball repo should copy (with minor edits) from the baseball repo:")
    add_bullet(doc, "ingestion/send-email.js — Resend API caller (change only the DEFAULT_FROM string and the snapshot path if different).")
    add_bullet(doc, "ingestion/lib/email-template.js — HTML builder (rebuild the template for pickleball content: tour standings, player of the week, tournament results, etc.).")
    add_bullet(doc, "Add the \"Send morning email\" step to pickleball's daily workflow YAML.")

    add_callout(
        doc,
        "Architectural note",
        "The design decision to use two separate GitHub repos (not one combined repo) was made explicitly — this preserves domain independence, clean URL paths, and per-project deploy history. The minor cost is duplicating the Resend API key as a secret in each repo.",
    )

    # ===== Part 6: GitHub vs Local =====
    add_heading(doc, "6.  GitHub vs Local — What Lives Where", level=1, color=NAVY)
    add_para(doc, "Short version: production email lives 100% in GitHub Actions. Local setup is for development and testing only.", italic=True, color=MUTED)

    add_table(
        doc,
        ["Concern", "GitHub (production)", "Local (dev / test only)"],
        [
            ["Where the job runs",
             "GitHub Actions runner (ubuntu-latest), fires at 07:00 UTC daily via cron in .github/workflows/daily.yml.",
             "Your Windows machine, when you run node ingestion/send-email.js manually."],
            ["Where RESEND_API_KEY lives",
             "GitHub repo Settings → Secrets. Never in files.",
             "Your local .env (gitignored). Only needed if you want to send real email from your laptop."],
            ["Where EMAIL_RECIPIENTS lives",
             "GitHub repo Settings → Secrets.",
             "Your local .env (or shell env), if testing. Usually just set to your own email for local tests."],
            ["How to disable in production",
             "Delete RESEND_API_KEY from GitHub Secrets. Ingestion continues; email is skipped.",
             "Set EMAIL_DRY_RUN=1 before running — prints the email without calling Resend."],
            ["How to edit the email template",
             "Edit ingestion/lib/email-template.js locally, commit, push. GitHub uses the committed version on next run.",
             "Edit + run node ingestion/send-email.js with EMAIL_DRY_RUN=1 to preview."],
        ],
        col_widths_in=[1.4, 2.6, 2.5],
    )

    add_heading(doc, "6.1  Local dry-run (no email sent)", level=2)
    add_code_block(doc,
        'cd "C:\\Users\\John & Cindy Gladden\\Desktop\\AI\\Claude\\Baseball Project"\n'
        "set RESEND_API_KEY=re_dummy\n"
        "set EMAIL_RECIPIENTS=test@example.com\n"
        "set EMAIL_DRY_RUN=1\n"
        "node ingestion/send-email.js"
    )
    add_para(doc, "Prints subject + plain-text preview + HTML byte count. Does NOT hit the Resend API. Safe to run anytime.")

    add_heading(doc, "6.2  Local live-send (actually emails)", level=2)
    add_para(doc, "Normally unnecessary — GitHub does this nightly. But useful after big template changes:")
    add_code_block(doc,
        "set RESEND_API_KEY=re_YOUR_REAL_KEY\n"
        "set EMAIL_RECIPIENTS=you@yourdomain.com\n"
        "node ingestion/send-email.js"
    )

    # ===== Part 7: Recipient management =====
    add_heading(doc, "7.  Recipient Management", level=1, color=NAVY)

    add_heading(doc, "7.1  Adding or removing people", level=2)
    add_numbered(doc, "Go to https://github.com/jjmgladden/baseball-daily/settings/secrets/actions.")
    add_numbered(doc, "Click Update next to EMAIL_RECIPIENTS.")
    add_numbered(doc, "Paste the full new list (GitHub only stores write-only; you cannot read the current value, so keep a personal copy if helpful).")
    add_numbered(doc, "Save. Next workflow run sends to the new list — no code change.")

    add_heading(doc, "7.2  The \"two parallel lists\" question (baseball + pickleball)", level=2)
    add_para(doc, "GitHub Secrets are scoped per repo. If you keep separate recipient lists for baseball and pickleball, you maintain them independently — adding a name requires editing the secret in both repos.")
    add_para(doc, "For a list that changes 1-3 times a year, manual sync is recommended. If it becomes painful, the three escalation options are:")

    add_table(
        doc,
        ["Option", "Effort", "Trade-off"],
        [
            ["A. Accept duplication (recommended)",
             "Zero. Edit in both repos when it changes.",
             "Cheap. Separate lists can even be an asset — different audiences care about different sports."],
            ["B. Shared private gist",
             "~5 min. Create gist, add fetch logic to both send-email.js scripts, add a GitHub PAT secret to both repos.",
             "Central list; adds a network dependency and an auth token to manage."],
            ["C. Shared private config repo",
             "~10 min. Repo jjmgladden/shared-config holds recipient files; both workflows check it out.",
             "Room for other shared config; more moving parts."],
            ["D. GitHub Organization Secrets",
             "~30 min + migration. Convert jjmgladden to an Organization; secrets shared across repos.",
             "Heavy migration (URL changes, repo moves). Overkill for two hobby projects."],
        ],
        col_widths_in=[2.3, 1.5, 3.0],
    )

    # ===== Part 8: Troubleshooting =====
    add_heading(doc, "8.  Troubleshooting", level=1, color=NAVY)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["Workflow succeeds but no email arrives",
             "First-ever email landed in spam.",
             "Mark not-spam once; future emails inbox-correctly."],
            ["Log: \"RESEND_API_KEY not set — skipping\"",
             "Secret name typo or not saved.",
             "Re-create the secret with exact name RESEND_API_KEY (case-sensitive)."],
            ["Log: Resend API 403",
             "Invalid API key.",
             "Regenerate key in Resend dashboard; update the GitHub Secret."],
            ["Log: Resend API 422 on sender domain",
             "EMAIL_FROM points to a domain not verified in Resend.",
             "Either drop EMAIL_FROM (falls back to onboarding@resend.dev) or complete domain verification."],
            ["Email arrives but CTA link loads old content",
             "Pages not yet rebuilt when email sent.",
             "Normal — Pages needs 1-2 min after the ingestion push. Content is at most 1-2 min stale."],
            ["Domain won't verify in Resend",
             "DNS records wrong, or Cloudflare proxy is on (orange cloud).",
             "Double-check record values; set Cloudflare to DNS-only (grey cloud) for those records."],
        ],
        col_widths_in=[2.1, 2.0, 2.7],
    )

    # ===== Part 9: Cost =====
    add_heading(doc, "9.  Costs — Full Summary", level=1, color=NAVY)
    add_table(
        doc,
        ["Item", "Cost", "Notes"],
        [
            ["Resend free tier",
             "$0/mo",
             "3,000 emails/month. Baseball + pickleball combined = ~150/mo at 3 recipients."],
            ["GitHub Actions minutes",
             "$0",
             "Public repos get unlimited free minutes."],
            ["GitHub Pages hosting",
             "$0",
             "Free for public repos."],
            ["Resend default sender (onboarding@resend.dev)",
             "$0",
             "Works fine; small disclaimer in footer."],
            ["Custom sender domain (optional)",
             "~$10/yr",
             "Registered at Cloudflare Registrar or similar. Eliminates disclaimer; brands the email."],
            ["Total ongoing",
             "$0 today, $10/yr if you add a custom domain",
             ""],
        ],
        col_widths_in=[2.5, 1.8, 2.7],
    )

    # ===== Part 10: Files =====
    add_heading(doc, "10.  Files Involved (Baseball Repo)", level=1, color=NAVY)
    add_table(
        doc,
        ["Path", "Role"],
        [
            [".github/workflows/daily.yml",
             "Adds \"Send morning email\" step after the snapshot commit. Reads RESEND_API_KEY, EMAIL_RECIPIENTS, EMAIL_FROM from Secrets."],
            ["ingestion/send-email.js",
             "Entry point. Reads latest.json, calls buildEmail(), POSTs to Resend. Graceful skip if credentials missing. Supports EMAIL_DRY_RUN=1."],
            ["ingestion/lib/email-template.js",
             "Builds subject / HTML / plain-text from the snapshot. Inline styles only for email-client compatibility."],
            ["ingestion/lib/env.js",
             "Tolerant .env loader (existing — reused for local dev runs)."],
            ["docs/email-setup.md",
             "Short public-facing setup guide (Markdown)."],
            ["docs/Daily-Email-Setup-Guide.docx",
             "This document — full design + setup record."],
            ["docs/knowledge-base.md § KB-0025",
             "Knowledge-base entry summarizing the feature."],
        ],
        col_widths_in=[2.6, 4.4],
    )

    # ===== Part 11: Key Q&A transcript =====
    add_heading(doc, "11.  Key Q&A From the Session (Verbatim)", level=1, color=NAVY)
    add_para(doc, "The following exchanges during the 2026-04-22 session clarified architecture decisions that aren't obvious from code alone. Captured verbatim for future reference.", italic=True, color=MUTED)

    add_heading(doc, "Q1. \"where is EMAIL_RECIPIENTS\"", level=3)
    add_para(doc, "Answer: It doesn't exist by default. You create it as a GitHub Secret at https://github.com/jjmgladden/baseball-daily/settings/secrets/actions, named exactly EMAIL_RECIPIENTS, value = comma-separated email addresses. GitHub Secrets are write-only — once saved, even you cannot read the value back on the web UI; only update or delete.")

    add_heading(doc, "Q2. \"if i am using the same email for both baseball and pickleball — which github repo does it reside in\"", level=3)
    add_para(doc, "Answer: Each repo has its own secrets — they don't travel across repos. However, the Resend account (and its API key) can be shared. The correct pattern is: one Resend account → paste the same RESEND_API_KEY into both repos as a secret. EMAIL_RECIPIENTS may legitimately differ between projects.")

    add_heading(doc, "Q3. \"am i keeping 2 parallel lists — one for bb and one for pb and i have to keep them synced up if i want them the same\"", level=3)
    add_para(doc, "Answer: Yes — with two separate repos, the lists are independent and any sync is manual. For a list that changes 1-3 times a year, this is not painful. The four escalation options if it becomes painful are documented in §7.2 above; the recommendation is to live with the duplication unless it actively bites you.")

    # ===== Part 12: Checklist =====
    add_heading(doc, "12.  Activation Checklist", level=1, color=NAVY)
    add_para(doc, "Print or copy this to track owner tasks:")

    baseball_checks = [
        "Sign up at https://resend.com/signup",
        "Create API key named \"Ozark Joe's Baseball Daily\", sending access",
        "Copy the re_... key to a scratch pad",
        "Add GitHub Secret RESEND_API_KEY to jjmgladden/baseball-daily",
        "Add GitHub Secret EMAIL_RECIPIENTS with comma-separated addresses",
        "Trigger workflow: gh workflow run daily.yml --repo jjmgladden/baseball-daily",
        "Confirm email arrives in inbox (check spam on first run)",
        "(Optional) Register custom domain and add EMAIL_FROM secret",
    ]
    add_heading(doc, "Baseball", level=3, color=CARDS_RED)
    for item in baseball_checks:
        add_bullet(doc, "☐  " + item)

    pickleball_checks = [
        "Reuse existing Resend account — do NOT create a second",
        "Reuse the same re_... API key",
        "Add GitHub Secret RESEND_API_KEY to jjmgladden/pickleball-daily (same value as baseball)",
        "Add GitHub Secret EMAIL_RECIPIENTS (pickleball-specific list, may differ from baseball)",
        "Copy send-email.js and email-template.js into pickleball repo; rebuild template for pickleball content",
        "Add \"Send morning email\" step to pickleball daily workflow YAML",
        "Trigger test: gh workflow run daily.yml --repo jjmgladden/pickleball-daily",
    ]
    add_heading(doc, "Pickleball (future)", level=3, color=GOLD)
    for item in pickleball_checks:
        add_bullet(doc, "☐  " + item)

    # ===== Footer =====
    add_heading(doc, "End of Document", level=2, color=MUTED)
    add_para(doc, "Ozark Joe's Baseball Daily Intelligence Report  ·  https://jjmgladden.github.io/baseball-daily/", italic=True, color=MUTED, size=10)

    out_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "docs",
        "Daily-Email-Setup-Guide.docx",
    )
    doc.save(out_path)
    print(f"[build-email-doc] wrote {out_path}")
    return out_path


if __name__ == "__main__":
    build()
